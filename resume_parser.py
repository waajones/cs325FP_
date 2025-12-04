# resume_parser.py - Resume file parsing implementation
# Extracts text from PDF, DOCX, and TXT resume files
# Implements IResumeParser interface (DIP)

import io
import logging
from typing import Optional
import re
from interfaces import IResumeParser

# Try to import PDF library (optional dependency)
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Try to import DOCX library (optional dependency)
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeParser(IResumeParser):
    
    # Initialize parser with logging
    def __init__(self):
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
    
    # Parse resume file based on extension (TXT, PDF, DOCX)
    def parse_file(self, file_path: str) -> Optional[str]:
        if not file_path:
            return None
        
        # Get file extension
        file_extension = file_path.lower().split('.')[-1]
        
        try:
            # Route to appropriate parser based on file type
            if file_extension == 'txt':
                return self._parse_txt_file(file_path)
            elif file_extension == 'pdf':
                return self._parse_pdf_file(file_path)
            elif file_extension in ['docx', 'doc']:
                return self._parse_docx_file(file_path)
            else:
                self.logger.error(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error parsing resume: {str(e)}")
            return None
    
    # Read plain text file with encoding fallback
    def _parse_txt_file(self, file_path: str) -> Optional[str]:
        try:
            # Try UTF-8 encoding first
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            self.logger.info("Successfully parsed TXT file")
            return text
        except UnicodeDecodeError:
            try:
                # Fall back to latin-1 encoding
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                self.logger.info("Successfully parsed TXT file with latin-1 encoding")
                return text
            except Exception as e:
                self.logger.error(f"Failed to decode TXT file: {str(e)}")
                return None
        except Exception as e:
            self.logger.error(f"Error parsing TXT file: {str(e)}")
            return None
    
    # Extract text from PDF file using PyPDF2
    def _parse_pdf_file(self, file_path: str) -> Optional[str]:
        if not PDF_AVAILABLE:
            self.logger.error("PyPDF2 library not available. Cannot parse PDF files.")
            return None
        
        try:
            import PyPDF2 as pdf_module
            
            # Open PDF file in binary mode
            with open(file_path, 'rb') as f:
                pdf_reader = pdf_module.PdfReader(f)
                
                # Extract text from each page
                text_parts = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        text_parts.append(page_text)
                
                # Combine all pages
                full_text = '\n'.join(text_parts)
                
                if full_text.strip():
                    self.logger.info(f"Successfully parsed PDF file with {len(pdf_reader.pages)} pages")
                    return full_text
                else:
                    self.logger.warning("PDF file appears to be empty or text extraction failed")
                    return None
                    
        except Exception as e:
            self.logger.error(f"Error parsing PDF file: {str(e)}")
            return None
    
    # Extract text from DOCX file using python-docx
    def _parse_docx_file(self, file_path: str) -> Optional[str]:
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx library not available. Cannot parse DOCX files.")
            return None
        
        try:
            import docx as docx_module
            
            # Open DOCX document
            doc = docx_module.Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            # Combine all text
            full_text = '\n'.join(text_parts)
            
            if full_text.strip():
                self.logger.info("Successfully parsed DOCX file")
                return full_text
            else:
                self.logger.warning("DOCX file appears to be empty")
                return None
                
        except Exception as e:
            self.logger.error(f"Error parsing DOCX file: {str(e)}")
            return None
    
    # Parse resume from uploaded file object (for web interfaces)
    def parse_resume(self, uploaded_file) -> Optional[str]:
        if not uploaded_file:
            return None
        
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        try:
            if file_extension == 'txt':
                return self._parse_txt(uploaded_file)
            elif file_extension == 'pdf':
                return self._parse_pdf(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                return self._parse_docx(uploaded_file)
            else:
                self.logger.error(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error parsing resume: {str(e)}")
            return None
    
    # Parse TXT from uploaded file object
    def _parse_txt(self, uploaded_file) -> Optional[str]:
        try:
            stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8"))
            text = stringio.read()
            self.logger.info("Successfully parsed TXT file")
            return text
        except UnicodeDecodeError:
            try:
                text = uploaded_file.getvalue().decode("latin-1")
                self.logger.info("Successfully parsed TXT file with latin-1 encoding")
                return text
            except Exception as e:
                self.logger.error(f"Failed to decode TXT file: {str(e)}")
                return None
        except Exception as e:
            self.logger.error(f"Error parsing TXT file: {str(e)}")
            return None
    
    # Parse PDF from uploaded file object
    def _parse_pdf(self, uploaded_file) -> Optional[str]:
        if not PDF_AVAILABLE:
            self.logger.error("PyPDF2 library not available. Cannot parse PDF files.")
            return None
        
        try:
            pdf_bytes = io.BytesIO(uploaded_file.getvalue())
            
            import PyPDF2 as pdf_module
            pdf_reader = pdf_module.PdfReader(pdf_bytes)
            
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                if page_text.strip():
                    text_parts.append(page_text)
            
            full_text = '\n'.join(text_parts)
            
            if full_text.strip():
                self.logger.info(f"Successfully parsed PDF file with {len(pdf_reader.pages)} pages")
                return full_text
            else:
                self.logger.warning("PDF file appears to be empty or text extraction failed")
                return None
                
        except Exception as e:
            self.logger.error(f"Error parsing PDF file: {str(e)}")
            return None
    
    # Parse DOCX from uploaded file object
    def _parse_docx(self, uploaded_file) -> Optional[str]:
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx library not available. Cannot parse DOCX files.")
            return None
        
        try:
            docx_bytes = io.BytesIO(uploaded_file.getvalue())
            
            import docx as docx_module
            doc = docx_module.Document(docx_bytes)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            
            if full_text.strip():
                self.logger.info("Successfully parsed DOCX file")
                return full_text
            else:
                self.logger.warning("DOCX file appears to be empty")
                return None
                
        except Exception as e:
            self.logger.error(f"Error parsing DOCX file: {str(e)}")
            return None
    
    # Check if extracted text looks like a valid resume
    def validate_resume_content(self, text: str) -> bool:
        if not text or len(text.strip()) < 50:
            return False
        
        text_lower = text.lower()
        
        # Common resume keywords to look for
        resume_keywords = [
            'experience', 'education', 'skills', 'work', 'employment',
            'university', 'college', 'degree', 'bachelor', 'master',
            'phone', 'email', 'address', 'linkedin', 'github',
            'summary', 'objective', 'qualifications', 'achievements'
        ]
        
        # Count how many resume keywords are present
        keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
        
        # Check for email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        has_email = bool(re.search(email_pattern, text))
        
        # Check for phone number pattern
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}'
        ]
        has_phone = any(re.search(pattern, text) for pattern in phone_patterns)
        
        # Score based on criteria met
        criteria_met = 0
        if keyword_count >= 3:
            criteria_met += 1
        if has_email:
            criteria_met += 1
        if has_phone:
            criteria_met += 1
        if len(text.split()) >= 100:
            criteria_met += 1
        
        # Valid if at least 2 criteria are met
        is_valid = criteria_met >= 2
        
        self.logger.info(f"Resume validation: {criteria_met}/4 criteria met, valid: {is_valid}")
        
        return is_valid
    
    # Extract contact information (email, phone, LinkedIn, GitHub) from resume
    def extract_contact_info(self, text: str) -> dict:
        contact_info = {
            'emails': [],
            'phones': [],
            'linkedin': None,
            'github': None
        }
        
        # Find email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        contact_info['emails'] = re.findall(email_pattern, text)
        
        # Find phone numbers
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            contact_info['phones'].extend(phones)
        
        # Find LinkedIn URL
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group()
        
        # Find GitHub URL
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info['github'] = github_match.group()
        
        return contact_info
